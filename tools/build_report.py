from pathlib import Path
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs" / "Informe_Proyecto_Bitacora_HCI.docx"
OUT.parent.mkdir(parents=True, exist_ok=True)

BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
INK = RGBColor(11, 37, 69)
MUTED = RGBColor(86, 97, 112)
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F4F6F9"
LIME = "D9FF43"
RED = "9B1C1C"


def set_run_font(run, name="Calibri", size=11, color=None, bold=None, italic=None):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa, indent=120):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl = table._tbl
    tbl_pr = tbl.tblPr

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent))
    tbl_ind.set(qn("w:type"), "dxa")

    grid = tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for index, cell in enumerate(row.cells):
            set_cell_width(cell, widths_dxa[index])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_table_borders(table, color="C9D1DC", size="6"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_repeat_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Página ")
    set_run_font(run, size=9, color=MUTED)
    field_begin = OxmlElement("w:fldChar")
    field_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    field_sep = OxmlElement("w:fldChar")
    field_sep.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    field_end = OxmlElement("w:fldChar")
    field_end.set(qn("w:fldCharType"), "end")
    run._r.extend([field_begin, instr, field_sep, text, field_end])


def configure_styles(doc):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ):
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for style_name in ("List Bullet", "List Number"):
        style = styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.188)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25


def add_para(doc, text="", *, bold_prefix=None, italic=False, color=None, align=None, space_after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.25
    if align is not None:
        p.alignment = align
    if bold_prefix and text.startswith(bold_prefix):
        r1 = p.add_run(bold_prefix)
        set_run_font(r1, bold=True, color=color or INK)
        r2 = p.add_run(text[len(bold_prefix):])
        set_run_font(r2, italic=italic, color=color or INK)
    else:
        r = p.add_run(text)
        set_run_font(r, italic=italic, color=color or INK)
    return p


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.375 + level * 0.25)
    p.paragraph_format.first_line_indent = Inches(-0.188)
    r = p.add_run(text)
    set_run_font(r)
    return p


def add_number(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.left_indent = Inches(0.375)
    p.paragraph_format.first_line_indent = Inches(-0.188)
    r = p.add_run(text)
    set_run_font(r)
    return p


def add_callout(doc, label, text, fill=LIGHT_GRAY):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [9360], indent=120)
    set_table_borders(table, color="D3DAE4", size="4")
    set_repeat_header(table.rows[0])
    cell = table.cell(0, 0)
    shade_cell(cell, fill)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(f"{label}  ")
    set_run_font(r, size=10, bold=True, color=DARK_BLUE)
    r2 = p.add_run(text)
    set_run_font(r2, size=10, color=INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def add_table(doc, headers, rows, widths_dxa, header_fill=LIGHT_BLUE, font_size=9.5):
    table = doc.add_table(rows=1, cols=len(headers))
    set_table_geometry(table, widths_dxa)
    set_table_borders(table)
    header = table.rows[0]
    set_repeat_header(header)
    for i, text in enumerate(headers):
        cell = header.cells[i]
        shade_cell(cell, header_fill)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(str(text))
        set_run_font(r, size=font_size, bold=True, color=INK)
    for row_values in rows:
        row = table.add_row()
        for i, text in enumerate(row_values):
            cell = row.cells[i]
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(str(text))
            set_run_font(r, size=font_size, color=INK)
    set_table_geometry(table, widths_dxa)
    return table


def add_caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(10)
    r = p.add_run(text)
    set_run_font(r, size=9, italic=True, color=MUTED)


def add_image(doc, relative_path, caption, width=6.2):
    image_path = ROOT / relative_path
    doc.add_picture(str(image_path), width=Inches(width))
    shape = doc.inline_shapes[-1]
    doc_pr = shape._inline.docPr
    doc_pr.set("title", caption)
    doc_pr.set("descr", caption)
    p = doc.paragraphs[-1]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_caption(doc, caption)


def add_header_footer(section):
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    header_p = section.header.paragraphs[0]
    header_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    header_p.paragraph_format.space_after = Pt(0)
    r = header_p.add_run("BITÁCORA DIGITAL · PROYECTO HCI")
    set_run_font(r, size=8.5, bold=True, color=MUTED)
    footer_p = section.footer.paragraphs[0]
    footer_p.paragraph_format.space_before = Pt(0)
    add_page_number(footer_p)


doc = Document()
configure_styles(doc)
section = doc.sections[0]
add_header_footer(section)

# Cover: editorial cover pattern adapted to a formal university report.
for _ in range(5):
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(18)
r = p.add_run("PROYECTO FINAL · HCI")
set_run_font(r, size=12, bold=True, color=BLUE)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(8)
r = p.add_run("Bitácora Digital")
set_run_font(r, size=30, bold=True, color=INK)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(25)
r = p.add_run("Portafolio para organizar, consultar y evaluar evidencias del curso")
set_run_font(r, size=15, color=MUTED)
add_callout(
    doc,
    "Propósito",
    "Documentar el problema, las decisiones de HCI, el proceso de diseño y la implementación de una interfaz centrada en estudiantes.",
    fill="F1F7E2",
)
for _ in range(4):
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Equipo de desarrollo · Julio 2026")
set_run_font(r, size=11, color=MUTED)
doc.add_page_break()

add_para(doc, "Este informe acompaña el sistema publicado en GitHub y organiza la evidencia solicitada en el documento de requisitos del curso.", italic=True, color=MUTED)

doc.add_heading("Resumen ejecutivo", level=1)
add_para(doc, "Bitácora Digital es un portafolio web para estudiantes que necesitan reunir actividades, archivos y evidencias de HCI en un solo lugar. La solución prioriza tres acciones: registrar una evidencia, encontrarla por categoría o búsqueda y revisarla en una vista de evaluación.")
add_para(doc, "La propuesta se implementó como una aplicación web estática con HTML, CSS y JavaScript. Las evidencias descriptivas se conservan en el navegador mediante localStorage; los adjuntos pueden enviarse a Cloudinary y, de forma opcional, importarse desde la carpeta principal de OneDrive mediante Microsoft Graph.")
add_callout(doc, "Estado del proyecto", "El CRUD obligatorio, la navegación, la validación, la búsqueda, los filtros, la vista de evaluación, el tema claro y la adaptación móvil fueron verificados con una prueba automatizada local. La prueba con participantes reales debe ser ejecutada por el grupo antes de la entrega.")

doc.add_heading("1. Problema y oportunidad", level=1)
add_para(doc, "Las evidencias de un curso suelen quedar repartidas entre carpetas locales, servicios de nube y conversaciones. Esa dispersión dificulta saber qué se entregó, a qué tipo pertenece cada actividad y cómo mostrar el avance durante una revisión.")
add_para(doc, "El proyecto propone una bitácora digital con lenguaje cercano al curso: talleres, laboratorios, parciales y proyectos. La oportunidad de diseño consiste en convertir un conjunto de archivos en una experiencia de seguimiento comprensible, consultable y fácil de demostrar.")

doc.add_heading("2. Objetivos", level=1)
doc.add_heading("2.1 Objetivo general", level=2)
add_para(doc, "Diseñar y desarrollar un portafolio digital que permita organizar, visualizar y revisar evidencias del curso de HCI aplicando principios de Interacción Humano-Computador y Diseño Centrado en el Usuario.")
doc.add_heading("2.2 Objetivos específicos", level=2)
for item in (
    "Centralizar el registro de actividades y sus archivos adjuntos.",
    "Organizar las evidencias en talleres, laboratorios, parciales y proyectos.",
    "Reducir el tiempo de búsqueda mediante filtros, búsqueda global y navegación por categorías.",
    "Prevenir errores con validación de campos, límites de archivo y confirmación antes de eliminar.",
    "Ofrecer una vista de evaluación que facilite la revisión por parte del estudiante o docente.",
    "Documentar el proceso de diseño, las decisiones de interfaz y las limitaciones pendientes.",
):
    add_bullet(doc, item)

doc.add_heading("3. Usuarios y contexto de uso", level=1)
add_para(doc, "El usuario primario es un estudiante universitario que registra entregas durante el semestre. El usuario secundario es un compañero o docente que necesita revisar rápidamente títulos, categorías, fechas y archivos.")
add_image(doc, "evidencias/wireframes/png/01-dashboard.png", "Figura 1. Proto-arquitectura del resumen: iniciar, registrar y consultar estado.", width=6.2)
doc.add_heading("3.1 Proto-persona", level=2)
add_para(doc, "Alex representa una hipótesis de usuario derivada del contexto del curso: trabaja con entregas distribuidas, necesita nombres claros y quiere comprobar que una acción quedó guardada. Este perfil debe validarse con estudiantes reales; no se presenta como una entrevista realizada.")
add_table(doc, ["Necesidad", "Respuesta de diseño"], [
    ("Encontrar una evidencia", "Búsqueda global, filtros y acceso rápido por categoría."),
    ("Saber qué ocurrió", "Toasts, conteos, fechas y mensajes de error."),
    ("Revisar el avance", "Vista de evaluación agrupada por tipo."),
    ("Evitar pérdida de información", "Confirmación al eliminar y persistencia local."),
], [1800, 7560], font_size=9.5)

doc.add_heading("4. Requisitos funcionales y no funcionales", level=1)
add_table(doc, ["Requisito", "Implementación", "Estado"], [
    ("Registrar evidencia", "Modal con título, descripción, categoría y archivo opcional.", "Cumple"),
    ("Organizar contenido", "Talleres, Laboratorios, Parciales y Proyectos.", "Cumple"),
    ("Visualizar contenido", "Tarjetas de archivos y tablas de evaluación.", "Cumple"),
    ("Editar y eliminar", "Formulario de edición y confirmación de eliminación.", "Cumple"),
    ("Buscar y filtrar", "Búsqueda global y filtros por categoría.", "Cumple"),
    ("Interfaz usable", "Jerarquía clara, feedback, teclado y responsive.", "Cumple"),
    ("Nube", "Cloudinary para adjuntos y OneDrive opcional con Microsoft Graph.", "Parcial/configurable"),
], [2600, 5600, 1160], font_size=9)
add_para(doc, "No se requiere una base de datos externa para el alcance actual: localStorage resuelve la persistencia del prototipo. Si el sistema se convierte en un producto multiusuario, deberá migrarse a una API y una base de datos con autenticación.", italic=True, color=MUTED)

doc.add_heading("5. Proceso de HCI y DCU", level=1)
add_para(doc, "El proceso se documentó como una cadena de decisiones: interpretar el brief, formular una hipótesis de usuario, organizar la información, revisar la interfaz con heurísticas y recorrer las tareas principales. Los artefactos reproducibles se encuentran en `evidencias/dcu/`.")
doc.add_heading("5.1 Técnicas aplicadas", level=2)
add_table(doc, ["Técnica", "Cómo se ejecutó", "Resultado"], [
    ("Matriz de requisitos", "Se comparó cada requisito del brief con archivos, funciones y capturas.", "Trazabilidad del alcance."),
    ("Proto-persona", "Se modeló un estudiante como hipótesis explícita, sin atribuirle entrevistas inventadas.", "Necesidades y frustraciones a validar."),
    ("Arquitectura de información", "Se ordenaron resumen, archivos, categorías y evaluación.", "Flujo de navegación y wireframes."),
    ("Evaluación heurística", "Se recorrieron crear, buscar, editar, mover, duplicar, eliminar y evaluar.", "Correcciones de idioma, ayuda y acciones."),
    ("Recorrido cognitivo", "Se simuló el primer registro de un laboratorio con un PDF.", "Puntos de orientación y feedback."),
    ("Prueba con usuarios", "Se preparó un guion, hoja de observación y plantilla de resultados.", "Pendiente de ejecución por el grupo."),
], [2000, 5000, 2360], font_size=8.8)
doc.add_heading("5.2 Evaluación heurística y mejoras", level=2)
add_para(doc, "La revisión encontró una mezcla de español e inglés, acciones que solo mostraban un mensaje de próxima versión y falta de ayuda inicial. Se corrigieron esos puntos: se tradujeron las acciones, mover cambia la categoría, duplicar crea una nueva evidencia y se agregó un modal de ayuda.")
add_image(doc, "evidencias/wireframes/png/02-archivos.png", "Figura 2. Wireframe de archivos: filtrar, seleccionar y administrar.", width=6.2)
add_heading = doc.add_heading
add_heading("5.3 Prueba de usabilidad pendiente", level=2)
add_para(doc, "El grupo debe aplicar el instrumento a cinco estudiantes, registrar tiempos y errores y actualizar `evidencias/dcu/07-plantilla-resultados.md`. Esto no debe rellenarse con datos inventados. La plantilla ya contiene tareas, preguntas de satisfacción y una tabla de observación.")

doc.add_heading("6. Diseño de la interfaz", level=1)
add_para(doc, "La interfaz usa un tema oscuro moderno como vista principal y un tema claro clásico para la lista de archivos. La barra lateral mantiene el contexto; la cabecera concentra búsqueda y ayuda; el contenido cambia entre resumen, archivos y evaluación.")
add_image(doc, "evidencias/capturas/01-resumen.png", "Figura 3. Captura verificada del resumen en tema oscuro.", width=6.2)
add_image(doc, "evidencias/capturas/03-archivos.png", "Figura 4. Captura verificada de tarjetas, categorías y acciones.", width=6.2)
add_image(doc, "evidencias/capturas/05-tema-claro.png", "Figura 5. Captura verificada del tema claro clásico.", width=6.2)
doc.add_heading("6.1 Wireframes", level=2)
add_image(doc, "evidencias/wireframes/png/03-evaluacion.png", "Figura 6. Wireframe de la vista de evaluación agrupada.", width=6.2)
add_para(doc, "Los wireframes priorizan jerarquía y flujo antes de los detalles visuales: el resumen responde “¿por dónde empiezo?”, archivos responde “¿dónde está y qué puedo hacer?” y evaluación responde “¿qué entregas debo revisar?”.")

doc.add_heading("7. Implementación técnica", level=1)
add_table(doc, ["Capa", "Decisión", "Beneficio"], [
    ("Estructura", "HTML semántico, modales, navegación y SVG inline.", "Sin imágenes externas para la iconografía."),
    ("Estilo", "CSS Grid/Flexbox, variables de tema y media queries.", "Responsive y cambio de tema consistente."),
    ("Estado", "JavaScript modular en un IIFE.", "Reduce variables globales y mantiene el flujo controlado."),
    ("Persistencia", "localStorage con migración de categorías antiguas.", "El prototipo funciona sin backend."),
    ("Archivos", "Cloudinary para adjuntos; validación de tipo y tamaño.", "Evita guardar bytes pesados en localStorage."),
    ("OneDrive", "MSAL Browser + Microsoft Graph con permisos explícitos.", "Importación opcional de archivos de Microsoft 365."),
], [1600, 5100, 2660], font_size=8.8)
add_para(doc, "Configuración de nube: OneDrive requiere una aplicación SPA en Microsoft Entra, una URI de redirección local y los permisos delegados `User.Read` y `Files.Read`. Sin `clientId`, la función no bloquea el resto de la aplicación.", italic=True, color=MUTED)

doc.add_heading("8. Validación y evidencias de funcionamiento", level=1)
add_para(doc, "Se creó `tests/ui-smoke.mjs`, una prueba automatizada que abre la aplicación, limpia el estado de prueba y recorre validación, creación, búsqueda, duplicación, movimiento, evaluación, cambio de tema y vista móvil. La prueba terminó correctamente el 27 de julio de 2026.")
add_table(doc, ["Recorrido", "Resultado observado"], [
    ("Validación", "Los tres campos obligatorios muestran error cuando se envían vacíos."),
    ("CRUD", "La evidencia se crea, se edita, se duplica y se elimina con confirmación."),
    ("Búsqueda", "La consulta “accesibilidad” devuelve dos registros relacionados."),
    ("Evaluación", "Las entregas aparecen agrupadas por categoría en tablas."),
    ("Tema y responsive", "El tema claro y la vista móvil se capturan sin errores JavaScript."),
], [2600, 6760], font_size=9)
add_image(doc, "evidencias/capturas/04-evaluacion.png", "Figura 7. Captura verificada de la vista de evaluación.", width=6.2)
add_image(doc, "evidencias/capturas/06-vista-movil.png", "Figura 8. Captura verificada de la adaptación móvil.", width=3.0)

doc.add_heading("9. Limitaciones y trabajo futuro", level=1)
for item in (
    "La información descriptiva vive en el navegador y no se sincroniza entre dispositivos.",
    "Cloudinary usa una configuración de carga no firmada: para producción debe protegerse mediante un backend.",
    "La importación de OneDrive requiere configurar Microsoft Entra y permisos del propietario de la cuenta.",
    "La prueba con participantes reales aún debe ejecutarse y sus resultados deben reemplazar la plantilla.",
    "No se implementó colaboración multiusuario ni invitación de miembros porque excede el alcance mínimo del brief.",
):
    add_bullet(doc, item)

doc.add_heading("10. Entregables del curso", level=1)
add_table(doc, ["Entregable", "Ubicación", "Estado"], [
    ("Sistema funcional", "`index.html`, `styles.css`, `app.js`", "Listo"),
    ("Código fuente", "Repositorio GitHub", "Listo en rama de trabajo"),
    ("Documento del proyecto", "Este informe", "Listo"),
    ("Wireframes", "`evidencias/wireframes/`", "Listo"),
    ("Evidencias DCU", "`evidencias/dcu/`", "Listo + prueba pendiente"),
    ("Capturas", "`evidencias/capturas/`", "Listo"),
    ("Presentación", "`docs/Presentacion_Final_Bitacora_HCI.pptx`", "En preparación"),
], [2600, 5000, 1760], font_size=9)

doc.add_heading("11. Conclusiones", level=1)
add_para(doc, "Bitácora Digital transforma un conjunto disperso de entregas en un portafolio con una estructura que el estudiante puede reconocer: resumen, archivos y evaluación. La solución aplica visibilidad del estado, prevención de errores, consistencia visual, navegación por categorías y retroalimentación inmediata.")
add_para(doc, "El proyecto ya cuenta con un sistema funcional y evidencia técnica reproducible. Para cerrar la entrega académica, el grupo debe completar la prueba con estudiantes, adjuntar sus resultados reales, reemplazar los nombres del equipo si el docente los exige y presentar el flujo completo en la exposición.")

doc.add_heading("Fuentes y trazabilidad", level=1)
add_para(doc, "1. Documento del curso: `PROYECTO FINAL DEL CURSO HCI.pdf`, incluido en la carpeta local del proyecto.")
add_para(doc, "2. Repositorio: `https://github.com/devG3r4/HCI-PytFinal`.")
add_para(doc, "3. Evidencias internas: `evidencias/dcu/`, `evidencias/wireframes/`, `evidencias/capturas/` y `tests/ui-smoke.mjs`.")

doc.save(OUT)
print(OUT)
